#!/usr/bin/env python3
"""
Create analysis template from existing analysis file
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

def extract_analysis_structure():
    """Extract structure from existing analysis file"""
    analysis_dir = Path("dentist_consult_skill_analysis")
    
    # Find an existing analysis file
    analysis_files = list(analysis_dir.glob("*.docx"))
    if not analysis_files:
        print("❌ No existing analysis files found")
        return None
    
    # Use the first analysis file
    existing_file = analysis_files[0]
    print(f"📄 Analyzing: {existing_file.name}")
    
    try:
        doc = Document(existing_file)
        
        print("📋 Document Structure:")
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                if paragraph.style.name.startswith('Heading'):
                    print(f"  {paragraph.style.name}: {paragraph.text}")
                else:
                    print(f"  Content: {paragraph.text[:50]}...")
        
        return doc
        
    except Exception as e:
        print(f"❌ Error reading DOCX: {e}")
        return None

def create_analysis_template():
    """Create analysis template based on existing structure"""
    doc = Document()
    
    # Set up styles
    styles = doc.styles
    
    # Title style
    title_style = styles['Title']
    title_font = title_style.font
    title_font.name = 'Calibri'
    title_font.size = Pt(18)
    title_font.bold = True
    
    # Heading 1 style
    heading1_style = styles['Heading 1']
    heading1_font = heading1_style.font
    heading1_font.name = 'Calibri'
    heading1_font.size = Pt(14)
    heading1_font.bold = True
    
    # Normal style
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(11)
    
    # Add title
    title = doc.add_heading('Consultation Skill Analysis', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add analysis structure based on existing files
    doc.add_heading('Overall Assessment', level=1)
    doc.add_paragraph('[Brief but comprehensive assessment of the consultation quality, including professionalism, effectiveness, and patient engagement. Be critical. The dentist likes critical feedback more than positive feedback.]')
    
    doc.add_heading('Strengths', level=1)
    doc.add_paragraph('- [Strength 1]')
    doc.add_paragraph('- [Strength 2]')
    doc.add_paragraph('- [Strength 3]')
    
    doc.add_heading('Areas for Improvement', level=1)
    doc.add_paragraph('- [Area 1]')
    doc.add_paragraph('- [Area 2]')
    doc.add_paragraph('- [Area 3]')
    
    doc.add_heading('Communication Skills', level=1)
    doc.add_paragraph('- [Communication observation 1]')
    doc.add_paragraph('- [Communication observation 2]')
    
    doc.add_heading('Clinical Presentation', level=1)
    doc.add_paragraph('- [Clinical presentation observation 1]')
    doc.add_paragraph('- [Clinical presentation observation 2]')
    
    doc.add_heading('Financial Communication Analysis', level=1)
    doc.add_paragraph('- [Financial communication observation 1]')
    doc.add_paragraph('- [Financial communication observation 2]')
    
    doc.add_heading('Medical Information Handling', level=1)
    doc.add_paragraph('- [Medical information observation 1]')
    doc.add_paragraph('- [Medical information observation 2]')
    
    doc.add_heading('Medication Notes', level=1)
    doc.add_paragraph('- [Medication discussion observation 1]')
    doc.add_paragraph('- [Medication discussion observation 2]')
    doc.add_paragraph('- [Medication safety considerations]')
    
    doc.add_heading('Recommendations', level=1)
    doc.add_paragraph('- [Recommendation 1]')
    doc.add_paragraph('- [Recommendation 2]')
    doc.add_paragraph('- [Recommendation 3]')
    
    doc.add_heading('Key Takeaways', level=1)
    doc.add_paragraph('[Summary of main points and actionable insights]')
    
    # Save template
    template_path = Path("analysis_template.docx")
    doc.save(str(template_path))
    print(f"✅ Analysis template created: {template_path}")
    
    return template_path

def create_markdown_template():
    """Create analysis template in Markdown format"""
    md_content = """# Consultation Skill Analysis

## Overall Assessment

[Brief but comprehensive assessment of the consultation quality, including professionalism, effectiveness, and patient engagement. Be critical. The dentist likes critical feedback more than positive feedback.]

## Strengths

- [Strength 1]
- [Strength 2]
- [Strength 3]

## Areas for Improvement

- [Area 1]
- [Area 2]
- [Area 3]

## Communication Skills

- [Communication observation 1]
- [Communication observation 2]

## Clinical Presentation

- [Clinical presentation observation 1]
- [Clinical presentation observation 2]

## Financial Communication Analysis

- [Financial communication observation 1]
- [Financial communication observation 2]

## Medical Information Handling

- [Medical information observation 1]
- [Medical information observation 2]

## Medication Notes

- [Medication discussion observation 1]
- [Medication discussion observation 2]
- [Medication safety considerations]

## Recommendations

- [Recommendation 1]
- [Recommendation 2]
- [Recommendation 3]

## Key Takeaways

[Summary of main points and actionable insights]
"""
    
    # Save template
    template_path = Path("analysis_template.md")
    with open(template_path, 'w', encoding='utf-8') as f:
        f.write(md_content)
    print(f"✅ Markdown template created: {template_path}")
    
    return template_path

def main():
    """Main function"""
    print("📄 Creating Analysis Templates")
    print("=" * 40)
    
    # First extract structure from existing file
    extract_analysis_structure()
    
    # Then create both templates
    create_analysis_template()
    create_markdown_template()
    
    print("\n✅ Analysis templates created successfully!")

if __name__ == "__main__":
    main() 