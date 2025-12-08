#!/usr/bin/env python3
"""
Create summary templates for .md and .docx formats
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

def extract_summary_structure():
    """Extract structure from existing summary file"""
    summaries_dir = Path("summaries")
    
    # Find an existing summary file
    summary_files = list(summaries_dir.glob("*_Clean_Consultation_Summary.docx"))
    if not summary_files:
        print("❌ No existing summary files found")
        return None
    
    # Use the first summary file
    existing_file = summary_files[0]
    print(f"📄 Analyzing: {existing_file.name}")
    
    try:
        doc = Document(existing_file)
        
        print("📋 Document Structure:")
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                if paragraph.style.name.startswith('Heading'):
                    print(f"  {paragraph.style.name}: {paragraph.text}")
                else:
                    print(f"  Content: {paragraph.text[:50]}...")
        
        return doc
        
    except Exception as e:
        print(f"❌ Error reading DOCX: {e}")
        return None

def create_summary_template():
    """Create summary template based on existing structure"""
    doc = Document()
    
    # Set up styles
    styles = doc.styles
    
    # Title style
    title_style = styles['Title']
    title_font = title_style.font
    title_font.name = 'Calibri'
    title_font.size = Pt(18)
    title_font.bold = True
    
    # Heading 1 style
    heading1_style = styles['Heading 1']
    heading1_font = heading1_style.font
    heading1_font.name = 'Calibri'
    heading1_font.size = Pt(14)
    heading1_font.bold = True
    
    # Normal style
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(11)
    
    # Add title
    title = doc.add_heading('Consultation Summary', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add summary structure based on existing files
    doc.add_heading('Consultation Overview', level=1)
    doc.add_paragraph('- Date and type of consultation')
    doc.add_paragraph('- Main reason for visit')
    doc.add_paragraph('- Key topics discussed')
    
    doc.add_heading('Patient Information', level=1)
    doc.add_paragraph('- Current dental concerns')
    doc.add_paragraph('- Medical history mentioned')
    doc.add_paragraph('- Current medications (if discussed)')
    
    doc.add_heading('Treatment Discussion', level=1)
    doc.add_paragraph('- Treatment options presented')
    doc.add_paragraph('- Recommendations made')
    doc.add_paragraph('- Treatment plan outlined')
    
    doc.add_heading('Financial Information', level=1)
    doc.add_paragraph('- Costs discussed')
    doc.add_paragraph('- Insurance coverage mentioned')
    doc.add_paragraph('- Payment arrangements')
    
    doc.add_heading('Medication Information', level=1)
    doc.add_paragraph('- Medications discussed')
    doc.add_paragraph('- Medication history')
    doc.add_paragraph('- Medication safety considerations')
    
    doc.add_heading('Follow-up', level=1)
    doc.add_paragraph('- Next steps recommended')
    doc.add_paragraph('- Appointments scheduled')
    doc.add_paragraph('- Instructions given to patient')
    
    doc.add_heading('Key Points', level=1)
    doc.add_paragraph('- Important decisions made')
    doc.add_paragraph('- Critical information shared')
    doc.add_paragraph('- Action items for patient')
    
    # Save template
    template_path = Path("summary_template.docx")
    doc.save(str(template_path))
    print(f"✅ Summary template created: {template_path}")
    
    return template_path

def create_markdown_template():
    """Create summary template in Markdown format"""
    md_content = """# Consultation Summary

## Consultation Overview

- Date and type of consultation
- Main reason for visit
- Key topics discussed

## Patient Information

- Current dental concerns
- Medical history mentioned
- Current medications (if discussed)

## Treatment Discussion

- Treatment options presented
- Recommendations made
- Treatment plan outlined

## Financial Information

- Costs discussed
- Insurance coverage mentioned
- Payment arrangements

## Medication Information

- Medications discussed
- Medication history
- Medication safety considerations

## Follow-up

- Next steps recommended
- Appointments scheduled
- Instructions given to patient

## Key Points

- Important decisions made
- Critical information shared
- Action items for patient
"""
    
    # Save template
    template_path = Path("summary_template.md")
    with open(template_path, 'w', encoding='utf-8') as f:
        f.write(md_content)
    print(f"✅ Markdown template created: {template_path}")
    
    return template_path

def main():
    """Main function"""
    print("📄 Creating Summary Templates")
    print("=" * 40)
    
    # First extract structure from existing file
    extract_summary_structure()
    
    # Then create both templates
    create_summary_template()
    create_markdown_template()
    
    print("\n✅ Summary templates created successfully!")

if __name__ == "__main__":
    main() 